import os
import re
from typing import Annotated, Optional, TypedDict, List, Dict, Any
from dotenv import load_dotenv
from pydantic import BaseModel, Field

load_dotenv()

# Optional dependencies handling
try:
    from langchain_core.messages import AIMessage, HumanMessage
    from langchain_core.prompts import ChatPromptTemplate
    from langgraph.graph import END, START, StateGraph
    from langgraph.graph.message import add_messages
except ImportError:
    AIMessage = None
    HumanMessage = None
    ChatPromptTemplate = None
    END = "END"
    START = "START"
    StateGraph = None
    add_messages = None

try:
    from langchain_groq import ChatGroq
except ImportError:
    ChatGroq = None

try:
    from langchain_tavily import TavilySearch
except ImportError:
    TavilySearch = None

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    Document = None
    Pt = None

try:
    from reportlab.platypus import Paragraph, SimpleDocTemplate
    from reportlab.lib.styles import getSampleStyleSheet
except ImportError:
    Paragraph = None
    SimpleDocTemplate = None
    getSampleStyleSheet = None


def get_llm():
    if ChatGroq is None:
        return None
    try:
        groq_api_key = os.getenv("GROQ_API_KEY")
        if not groq_api_key:
            return None
        return ChatGroq(model="llama-3.3-70b-versatile", temperature=0.3, groq_api_key=groq_api_key)
    except Exception:
        return None


llm = get_llm()


class UserProfile(BaseModel):
    age: Optional[int] = Field(default=None)
    gender: Optional[str] = Field(default=None)
    weight: Optional[float] = Field(default=None)
    height: Optional[float] = Field(default=None)
    sport: Optional[str] = Field(default=None)
    goal: Optional[str] = Field(default=None)
    activity: Optional[str] = Field(default=None)
    level: Optional[str] = Field(default=None)
    training_days: Optional[int] = Field(default=None)
    lift_weight: Optional[float] = Field(default=None)
    reps: Optional[int] = Field(default=None)


EXTRACT_PROMPT = (
    ChatPromptTemplate.from_messages(
        [
            (
                "system",
                """
Extract the following information if available.
Return ONLY structured data.

Age
Gender
Weight
Height
Sport
Goal
Activity Level
Training Level
Training Days
Lift Weight
Repetitions

If information is missing, leave it null.
""",
            ),
            ("human", "{query}"),
        ]
    )
    if ChatPromptTemplate
    else None
)

PHYSIO_PROMPT = (
    ChatPromptTemplate.from_messages(
        [
            (
                "system",
                """
You are an experienced sports physiotherapist.
Use the provided tool results to answer the user's question.
Explain the injury, suggest appropriate stretches, mention recovery time, and note any red flags.
""",
            ),
            (
                "human",
                """
User Query:
{query}

Tool Results:
Symptom: {symptom}
Severity: {severity}
Recommended Stretches: {recommended_stretches}
Recovery Time: {recovery_time}
Prevention Tips: {prevention_tips}
Emergency: {emergency}
""",
            ),
        ]
    )
    if ChatPromptTemplate
    else None
)

TRAINING_PROMPT = (
    ChatPromptTemplate.from_messages(
        [
            (
                "system",
                """
You are an experienced sports training coach.
Use the provided user details and fitness metrics to recommend a training plan, recovery strategy, and performance tips.
""",
            ),
            (
                "human",
                """
User Query:
{query}

User Data:
{user_data}

BMI: {bmi}
Estimated 1RM: {one_rm}
Daily Calories: {calories}
Weekly Schedule: {schedule}
Progressive Overload Recommendation: {overload}
Recovery Advice: {recovery}
Sport Focus: {focus}
Warmup Routine: {warmup}
Cooldown Routine: {cooldown}
""",
            ),
        ]
    )
    if ChatPromptTemplate
    else None
)

NUTRITION_PROMPT = (
    ChatPromptTemplate.from_messages(
        [
            (
                "system",
                """
You are a professional sports nutritionist.
Use the provided data to recommend calories, macronutrients, hydration, meal structure, and supplement guidance.
""",
            ),
            (
                "human",
                """
User Query:
{query}

User Data:
{user_data}

BMR: {bmr}
TDEE: {tdee}
Protein Requirement: {protein}
Carbohydrate Recommendation: {carbs}
Fat Recommendation: {fats}
Water Intake: {water}
Meal Plan: {meals}
Supplements: {supplements}
Healthy Tips: {healthy_tips}
""",
            ),
        ]
    )
    if ChatPromptTemplate
    else None
)

RESUME_PROMPT = (
    ChatPromptTemplate.from_messages(
        [
            (
                "system",
                """
You are an expert sports resume writer.
Create an ATS-friendly, professional resume using the provided details.
Keep the text factual and do not invent information.
""",
            ),
            (
                "human",
                """
Use the following resume details to write a professional sports resume:
{resume_text}
""",
            ),
        ]
    )
    if ChatPromptTemplate
    else None
)

GENERAL_PROMPT = (
    ChatPromptTemplate.from_messages(
        [
            (
                "system",
                """
You are a knowledgeable sports assistant.
Answer the user question clearly and accurately using the provided search results.
""",
            ),
            (
                "human",
                """
User Question:
{query}

Search Result:
{search_result}
""",
            ),
        ]
    )
    if ChatPromptTemplate
    else None
)


def _to_float(value):
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _normalize_text(value):
    if value is None:
        return ""
    return str(value).strip().lower()


def create_search_tool():
    if TavilySearch is None:
        return None
    tavily_api_key = os.getenv("TAVILY_API_KEY")
    try:
        if tavily_api_key:
            return TavilySearch(max_results=5, tavily_api_key=tavily_api_key)
        return TavilySearch(max_results=5)
    except Exception:
        return None


def web_search(query):
    try:
        search_tool = create_search_tool()
        if search_tool is None:
            return "Search unavailable: TavilySearch is not configured."
        return search_tool.invoke(query)
    except Exception as exc:
        return f"Search unavailable: {exc}"


def latest_news(sport):
    return web_search(f"Latest {sport} news")


def tournament_schedule(tournament):
    return web_search(f"{tournament} upcoming schedule")


def player_statistics(player):
    return web_search(f"{player} latest statistics")


def team_ranking(sport):
    return web_search(f"Latest {sport} rankings")


def sports_rules(sport):
    return web_search(f"{sport} official rules")


def sports_equipment(sport):
    return web_search(f"Essential equipment for {sport}")


def coaching_tips(sport):
    return web_search(f"{sport} coaching tips")


# Nutrition helpers
def calculate_bmr(weight, height, age, gender):
    weight = _to_float(weight)
    height = _to_float(height)
    age = _to_float(age)
    gender = _normalize_text(gender)
    if weight is None or height is None or age is None or not gender:
        return None
    if gender == "male":
        bmr = 10 * weight + 6.25 * height - 5 * age + 5
    else:
        bmr = 10 * weight + 6.25 * height - 5 * age - 161
    return round(bmr, 2)


def healthy_tips():
    return [
        "Eat enough protein.",
        "Stay hydrated.",
        "Eat fruits and vegetables.",
        "Sleep 7-9 hours.",
        "Avoid processed foods.",
    ]


activity_levels = {
    "sedentary": 1.2,
    "light": 1.375,
    "moderate": 1.55,
    "active": 1.725,
    "very active": 1.9,
}


def calculate_tdee(bmr, activity):
    bmr = _to_float(bmr)
    activity = _normalize_text(activity)
    if bmr is None:
        return None
    multiplier = activity_levels.get(activity, 1.55)
    return round(bmr * multiplier, 2)


def protein_requirement(weight, goal):
    weight = _to_float(weight)
    goal = _normalize_text(goal)
    if weight is None:
        return None
    if goal == "muscle gain":
        return round(weight * 2.2, 1)
    if goal == "fat loss":
        return round(weight * 2.0, 1)
    if goal == "maintenance":
        return round(weight * 1.6, 1)
    return round(weight * 1.8, 1)


def water_requirement(weight):
    weight = _to_float(weight)
    if weight is None:
        return None
    return round(weight * 35 / 1000, 2)


def carb_requirement(weight, sport):
    weight = _to_float(weight)
    sport = _normalize_text(sport)
    if weight is None:
        return None
    endurance = ["marathon", "cycling", "football", "running"]
    if sport in endurance:
        return round(weight * 6, 1)
    return round(weight * 4.5, 1)


def fat_requirement(calories):
    calories = _to_float(calories)
    if calories is None:
        return None
    fat_calories = calories * 0.25
    grams = fat_calories / 9
    return round(grams, 1)


meal_templates = {
    "muscle gain": {
        "Breakfast": "Oats + Eggs + Banana",
        "Lunch": "Rice + Chicken + Vegetables",
        "Snack": "Greek Yogurt + Nuts",
        "Dinner": "Fish + Sweet Potato + Salad",
    },
    "fat loss": {
        "Breakfast": "Oats + Egg Whites",
        "Lunch": "Grilled Chicken + Vegetables",
        "Snack": "Apple + Almonds",
        "Dinner": "Paneer/Chicken + Salad",
    },
    "maintenance": {
        "Breakfast": "Oats + Milk",
        "Lunch": "Rice + Dal + Vegetables",
        "Snack": "Fruit",
        "Dinner": "Chicken/Fish + Vegetables",
    },
}


def meal_plan(goal):
    goal = _normalize_text(goal)
    if goal in meal_templates:
        return meal_templates[goal]
    return meal_templates["maintenance"]


def supplements(goal):
    goal = _normalize_text(goal)
    if goal == "muscle gain":
        return ["Whey Protein", "Creatine", "Fish Oil", "Multivitamin"]
    if goal == "fat loss":
        return ["Protein Powder", "Fish Oil", "Electrolytes"]
    return ["Protein Powder", "Multivitamin"]


# Physio helpers
symptom_database = {
    "back pain": {"causes": ["Muscle strain", "Poor lifting technique", "Weak core muscles", "Overtraining"]},
    "knee pain": {"causes": ["Runner's Knee", "Patellar Tendonitis", "ACL Strain", "Meniscus Injury"]},
    "shoulder pain": {"causes": ["Rotator Cuff Injury", "Shoulder Impingement", "Muscle Tightness"]},
    "ankle pain": {"causes": ["Ankle Sprain", "Ligament Injury", "Overuse"]},
    "hamstring pain": {"causes": ["Hamstring Strain", "Muscle Tightness", "Overstretching"]},
    "neck pain": {"causes": ["Poor Posture", "Muscle Tightness", "Whiplash"]},
}


def symptom_checker(query: str):
    query = query.lower()
    for symptom in symptom_database:
        if symptom in query:
            return {"symptom": symptom, "causes": symptom_database[symptom]["causes"]}
    return {"symptom": "Unknown", "causes": ["Unable to identify the injury."]}


def severity_checker(query: str):
    query = query.lower()
    severe = ["fracture", "cannot walk", "unable to move", "dislocated", "broken bone", "severe swelling", "bleeding"]
    moderate = ["pain", "sprain", "strain", "tightness", "stiffness"]
    mild = ["sore", "slight pain", "minor discomfort"]
    for word in severe:
        if word in query:
            return "Severe"
    for word in moderate:
        if word in query:
            return "Moderate"
    for word in mild:
        if word in query:
            return "Mild"
    return "Unknown"


stretch_database = {
    "back pain": ["Cat-Cow Stretch", "Child's Pose", "Pelvic Tilt", "Hamstring Stretch"],
    "knee pain": ["Quadriceps Stretch", "Hamstring Stretch", "Calf Stretch"],
    "shoulder pain": ["Cross Body Stretch", "Doorway Stretch", "Pendulum Exercise"],
    "ankle pain": ["Calf Stretch", "Toe Raises", "Alphabet Exercise"],
    "hamstring pain": ["Standing Hamstring Stretch", "Seated Hamstring Stretch", "Foam Rolling"],
}


def recommend_stretches(query: str):
    query = query.lower()
    for injury in stretch_database:
        if injury in query:
            return stretch_database[injury]
    return ["No stretch recommendations found."]


def recovery_time(symptom: str):
    symptom = symptom.lower()
    recovery = {
        "back pain": "1-3 weeks",
        "knee pain": "2-6 weeks",
        "shoulder pain": "2-8 weeks",
        "ankle pain": "2-6 weeks",
        "hamstring pain": "2-6 weeks",
    }
    return recovery.get(symptom, "Depends on diagnosis")


def prevention_tips():
    return [
        "Warm up before training.",
        "Cool down after exercise.",
        "Maintain proper technique.",
        "Avoid sudden increases in training load.",
        "Stay hydrated.",
        "Sleep at least 7-9 hours.",
        "Include mobility exercises weekly.",
    ]


def emergency_check(severity):
    return severity == "Severe"


# Training helpers
def bmi_calculator(weight, height):
    weight = _to_float(weight)
    height = _to_float(height)
    if not weight or not height:
        return None
    height_m = height / 100
    bmi = weight / (height_m**2)
    if bmi < 18.5:
        category = "Underweight"
    elif bmi < 25:
        category = "Normal Weight"
    elif bmi < 30:
        category = "Overweight"
    else:
        category = "Obese"
    return {"BMI": round(bmi, 2), "Category": category}


def one_rep_max(weight, reps):
    weight = _to_float(weight)
    reps = _to_float(reps)
    if not weight or not reps:
        return None
    one_rm = weight * (1 + reps / 30)
    return round(one_rm, 2)


def calorie_calculator(weight, goal):
    weight = _to_float(weight)
    if not weight:
        return None
    goal = _normalize_text(goal)
    if goal == "weight loss" or goal == "fat loss":
        return round(weight * 25, 1)
    if goal == "maintenance":
        return round(weight * 30, 1)
    if goal == "muscle gain":
        return round(weight * 35, 1)
    return round(weight * 30, 1)


def weekly_schedule(level):
    level = _normalize_text(level)
    if level == "beginner":
        return {
            "Monday": "Full Body",
            "Tuesday": "Rest",
            "Wednesday": "Upper Body",
            "Thursday": "Rest",
            "Friday": "Lower Body",
            "Saturday": "Cardio",
            "Sunday": "Recovery",
        }
    if level == "intermediate":
        return {
            "Monday": "Push",
            "Tuesday": "Pull",
            "Wednesday": "Legs",
            "Thursday": "Cardio",
            "Friday": "Upper",
            "Saturday": "Lower",
            "Sunday": "Recovery",
        }
    return {
        "Monday": "Strength",
        "Tuesday": "Power",
        "Wednesday": "Mobility",
        "Thursday": "Speed",
        "Friday": "Strength",
        "Saturday": "Conditioning",
        "Sunday": "Recovery",
    }


def progressive_overload(weight):
    weight = _to_float(weight)
    if not weight:
        return None
    return round(weight * 1.025, 2)


def rest_recommendation(training_days):
    training_days = _to_float(training_days)
    if not training_days:
        return "1-2 Rest Days per week"
    if training_days <= 3:
        return "2-3 Rest Days"
    if training_days <= 5:
        return "1-2 Rest Days"
    return "Minimum 1 Full Recovery Day"


training_templates = {
    "football": {"focus": ["Sprint", "Agility", "Passing", "Endurance"]},
    "cricket": {"focus": ["Power", "Reaction", "Core", "Throwing"]},
    "basketball": {"focus": ["Vertical Jump", "Speed", "Balance", "Shooting"]},
    "badminton": {"focus": ["Footwork", "Reaction", "Agility", "Stamina"]},
}


def sport_training(sport):
    sport = _normalize_text(sport)
    if sport in training_templates:
        return training_templates[sport]["focus"]
    return ["Strength", "Cardio", "Mobility"]


def warmup():
    return ["5 min Jog", "Dynamic Stretching", "Jumping Jacks", "Mobility Exercises"]


def cooldown():
    return ["Walking", "Static Stretching", "Deep Breathing", "Foam Rolling"]


# Resume helpers
resume_prompt_template = (
    ChatPromptTemplate.from_messages(
        [
            (
                "system",
                """
You are an expert Sports Resume Writer.
Write a professional ATS friendly resume.
Improve grammar.
Improve wording.
Highlight achievements.
Do not invent information.
Use the following sections:
Career Objective
Sports Profile
Education
Experience
Achievements
Skills
Certificates
Languages
References (Optional)
Return clean formatted text.
""",
            ),
            (
                "human",
                """
Use the following resume details to write a professional sports resume:
{resume_text}
""",
            ),
        ]
    )
    if ChatPromptTemplate
    else None
)

ACTION_VERBS = [
    "Led",
    "Achieved",
    "Won",
    "Represented",
    "Organized",
    "Developed",
    "Improved",
    "Coached",
    "Managed",
    "Trained",
    "Mentored",
    "Participated",
    "Captained",
    "Qualified",
]


def generate_resume(llm_model, data):
    if llm_model and resume_prompt_template:
        chain = resume_prompt_template | llm_model
        resume_text = data.get("resume_text")
        if resume_text is None:
            resume_data = data.get("resume_data")
            if isinstance(resume_data, dict) and resume_data:
                resume_lines = []
                for key in [
                    "name",
                    "email",
                    "phone",
                    "location",
                    "sport",
                    "position",
                    "education",
                    "experience",
                    "achievements",
                    "skills",
                    "certificates",
                    "languages",
                ]:
                    value = resume_data.get(key)
                    if value:
                        resume_lines.append(f"{key.capitalize()}: {value}")
                resume_text = "\n".join(resume_lines)
        if resume_text is None:
            resume_text = data.get("query", "")
        if not resume_text:
            resume_text = "No resume details were provided. Please include your name, sport, experience, achievements, and skills."
        try:
            response = chain.invoke({"resume_text": resume_text})
            return response.content
        except Exception:
            pass

    rdata = data.get("resume_data", {}) or {}
    name = rdata.get("name", "Sports Athlete")
    sport = rdata.get("sport", "General Sports")
    lines = [
        f"=== {name.upper()} ===",
        f"Sport: {sport} | Location: {rdata.get('location', 'India')}",
        f"Email: {rdata.get('email', 'contact@playure.com')} | Phone: {rdata.get('phone', 'N/A')}",
        "\nCAREER OBJECTIVE:",
        f"Dedicated {sport} player seeking opportunities to excel in competitive leagues and represent top teams.",
        "\nSPORTS PROFILE:",
        f"• Discipline: {sport}",
        f"• Position/Role: {rdata.get('position', 'Player')}",
        "\nEXPERIENCE:",
        rdata.get("experience", "Active competitive player in regional and state level tournaments."),
        "\nACHIEVEMENTS:",
        rdata.get("achievements", "Participated in regional sports competitions with consistent performance records."),
        "\nSKILLS:",
        rdata.get("skills", "Team Leadership, Physical Conditioning, Tactical Analysis, High Performance Mindset."),
    ]
    return "\n".join(lines)


def ats_score(data):
    score = 0
    suggestions = []
    if data.get("name"):
        score += 2
    else:
        suggestions.append("Add your full name.")
    if data.get("email"):
        score += 2
    else:
        suggestions.append("Add your email address.")
    if data.get("phone"):
        score += 2
    else:
        suggestions.append("Add your phone number.")
    if data.get("location"):
        score += 2
    else:
        suggestions.append("Add your location.")
    if data.get("sport"):
        score += 2
    else:
        suggestions.append("Mention your sport.")

    objective = data.get("career_objective", "")
    if len(objective) > 40:
        score += 10
    else:
        suggestions.append("Write a stronger career objective.")

    profile = data.get("sports_profile", "")
    if len(profile) > 40:
        score += 10
    else:
        suggestions.append("Add a sports profile.")

    if data.get("education"):
        score += 10
    else:
        suggestions.append("Add education details.")

    experience = data.get("experience", "")
    if len(experience) > 50:
        score += 15
    else:
        suggestions.append("Describe your sports experience.")

    achievements = data.get("achievements", "")
    if len(achievements) > 40:
        score += 10
    else:
        suggestions.append("Mention sports achievements.")

    if re.search(r"\d", achievements):
        score += 5
    else:
        suggestions.append("Use measurable achievements (e.g. Won 3 tournaments).")

    skills = data.get("skills", "")
    if len(skills) > 20:
        score += 10
    else:
        suggestions.append("Mention technical and sports skills.")

    certificates = data.get("certificates", "")
    if certificates:
        score += 10
    else:
        suggestions.append("Include sports certifications.")

    resume_text = " ".join([objective, profile, experience, achievements])
    verb_found = any(verb.lower() in resume_text.lower() for verb in ACTION_VERBS)
    if verb_found:
        score += 5
    else:
        suggestions.append("Use action verbs like Led, Won, Achieved, Represented.")

    return score, suggestions


def generate_docx(name, resume):
    if Document is None or Pt is None:
        return None
    if not name:
        name = "sports_resume"
    safe_name = name.replace(" ", "_")
    document = Document()
    heading = document.add_heading(name, 0)
    if hasattr(heading, "style") and hasattr(heading.style, "font"):
        heading.style.font.size = Pt(22)
    for line in resume.split("\n"):
        if line.strip():
            document.add_paragraph(line)
    filename = f"{safe_name}.docx"
    document.save(filename)
    return filename


def generate_pdf(name, resume):
    if Paragraph is None or SimpleDocTemplate is None or getSampleStyleSheet is None:
        return None
    filename = f"{name}.pdf"
    doc = SimpleDocTemplate(filename)
    styles = getSampleStyleSheet()
    story = []
    for line in resume.split("\n"):
        if line.strip():
            story.append(Paragraph(line, styles["BodyText"]))
    doc.build(story)
    return filename


def build_resume(llm_model, data):
    resume = generate_resume(llm_model, data)
    resume_data = data.get("resume_data", {}) or {}
    ats, suggestions = ats_score(resume_data)
    name = resume_data.get("name") or "sports_resume"
    docx = generate_docx(name, resume)
    pdf = generate_pdf(name, resume)
    return {"resume": resume, "ats": ats, "suggestions": suggestions, "docx": docx, "pdf": pdf}


if llm is not None and ChatPromptTemplate is not None:
    try:
        structured_llm = llm.with_structured_output(UserProfile)
        extractor_chain = EXTRACT_PROMPT | structured_llm
    except Exception:
        extractor_chain = None
else:
    extractor_chain = None


def extract_information(query):
    if extractor_chain is None:
        return {
            "age": None,
            "gender": None,
            "weight": None,
            "height": None,
            "sport": None,
            "goal": None,
            "activity": None,
            "level": None,
            "training_days": None,
            "lift_weight": None,
            "reps": None,
        }
    try:
        result = extractor_chain.invoke({"query": query})
        return result.model_dump()
    except Exception:
        return {
            "age": None,
            "gender": None,
            "weight": None,
            "height": None,
            "sport": None,
            "goal": None,
            "activity": None,
            "level": None,
            "training_days": None,
            "lift_weight": None,
            "reps": None,
        }


class SportsState(TypedDict):
    messages: list
    category: str
    response: str
    user_data: dict
    resume_data: dict
    tool_output: dict


def physio_node(state):
    query = state["messages"][-1].content if hasattr(state["messages"][-1], "content") else str(state["messages"][-1])
    symptom = symptom_checker(query)
    severity = severity_checker(query)
    stretches = recommend_stretches(query)
    recovery = recovery_time(symptom["symptom"])
    prevention = prevention_tips()
    emergency = emergency_check(severity)

    if llm is None or PHYSIO_PROMPT is None:
        response_text = f"Physio guidance for '{query}':\n\nSymptom: {symptom['symptom'].capitalize()}\nSeverity: {severity}\nSuggested stretches: {', '.join(stretches)}\nRecovery time: {recovery}\nPrevention tips:\n• " + "\n• ".join(prevention)
        if emergency:
            response_text += "\n\n⚠️ WARNING: Severe condition detected. Please consult a medical professional immediately!"
        return {"messages": [AIMessage(content=response_text) if AIMessage else response_text], "response": response_text}

    try:
        response = (PHYSIO_PROMPT | llm).invoke(
            {
                "query": query,
                "symptom": symptom,
                "severity": severity,
                "recommended_stretches": stretches,
                "recovery_time": recovery,
                "prevention_tips": prevention,
                "emergency": emergency,
            }
        )
        return {"messages": [AIMessage(content=response.content) if AIMessage else response.content], "response": response.content}
    except Exception:
        response_text = f"Physio guidance for '{query}':\n\nSymptom: {symptom['symptom'].capitalize()}\nSeverity: {severity}\nSuggested stretches: {', '.join(stretches)}\nRecovery time: {recovery}\nPrevention tips:\n• " + "\n• ".join(prevention)
        return {"messages": [AIMessage(content=response_text) if AIMessage else response_text], "response": response_text}


def training_node(state):
    query = state["messages"][-1].content if hasattr(state["messages"][-1], "content") else str(state["messages"][-1])
    new_data = extract_information(query)
    data = {**state.get("user_data", {}), **{k: v for k, v in new_data.items() if v is not None}}

    bmi = None
    one_rm = None
    calories = None
    schedule = None
    overload = None
    recovery = None
    focus = None

    if data.get("weight") and data.get("height"):
        bmi = bmi_calculator(data["weight"], data["height"])
    if data.get("lift_weight") and data.get("reps"):
        one_rm = one_rep_max(data["lift_weight"], data["reps"])
    if data.get("weight") and data.get("goal"):
        calories = calorie_calculator(data["weight"], data["goal"])
    if data.get("level"):
        schedule = weekly_schedule(data["level"])
    if data.get("lift_weight"):
        overload = progressive_overload(data["lift_weight"])
    if data.get("training_days") is not None:
        recovery = rest_recommendation(data["training_days"])
    if data.get("sport"):
        focus = sport_training(data["sport"])

    if llm is None or TRAINING_PROMPT is None:
        response_text = f"Training guidance for '{query}':\n\nUser metrics: {data}\nBMI: {bmi}\nEstimated 1RM: {one_rm}\nWeekly Schedule: {schedule}\nFocus Areas: {focus}\nWarmup: {warmup()}\nCooldown: {cooldown()}"
        return {"messages": [AIMessage(content=response_text) if AIMessage else response_text], "response": response_text}

    try:
        response = (TRAINING_PROMPT | llm).invoke(
            {
                "query": query,
                "user_data": data,
                "bmi": bmi,
                "one_rm": one_rm,
                "calories": calories,
                "schedule": schedule,
                "overload": overload,
                "recovery": recovery,
                "focus": focus,
                "warmup": warmup(),
                "cooldown": cooldown(),
            }
        )
        return {"messages": [AIMessage(content=response.content) if AIMessage else response.content], "response": response.content}
    except Exception:
        response_text = f"Training guidance for '{query}':\n\nUser metrics: {data}\nBMI: {bmi}\nEstimated 1RM: {one_rm}\nWeekly Schedule: {schedule}"
        return {"messages": [AIMessage(content=response_text) if AIMessage else response_text], "response": response_text}


def nutrition_node(state):
    query = state["messages"][-1].content if hasattr(state["messages"][-1], "content") else str(state["messages"][-1])
    data = state.get("user_data", {})
    bmr = calculate_bmr(data.get("weight"), data.get("height"), data.get("age"), data.get("gender"))
    tdee = calculate_tdee(bmr, data.get("activity"))
    protein = protein_requirement(data.get("weight"), data.get("goal"))
    carbs = carb_requirement(data.get("weight"), data.get("sport"))
    fats = fat_requirement(tdee or 2000)
    water = water_requirement(data.get("weight"))
    meals = meal_plan(data.get("goal"))

    if llm is None or NUTRITION_PROMPT is None:
        response_text = f"Nutrition guidance for '{query}':\n\nBMR: {bmr} kcal\nTDEE: {tdee} kcal\nProtein requirement: {protein}g\nCarbs: {carbs}g\nFats: {fats}g\nWater intake: {water}L\nMeal Plan: {meals}\nSupplements: {supplements(data.get('goal'))}"
        return {"messages": [AIMessage(content=response_text) if AIMessage else response_text], "response": response_text}

    try:
        response = (NUTRITION_PROMPT | llm).invoke(
            {
                "query": query,
                "user_data": data,
                "bmr": bmr,
                "tdee": tdee,
                "protein": protein,
                "carbs": carbs,
                "fats": fats,
                "water": water,
                "meals": meals,
                "supplements": supplements(data.get("goal")),
                "healthy_tips": healthy_tips(),
            }
        )
        return {"messages": [AIMessage(content=response.content) if AIMessage else response.content], "response": response.content}
    except Exception:
        response_text = f"Nutrition guidance for '{query}':\n\nBMR: {bmr} kcal\nTDEE: {tdee} kcal\nProtein requirement: {protein}g\nCarbs: {carbs}g\nFats: {fats}g\nWater intake: {water}L\nMeal Plan: {meals}"
        return {"messages": [AIMessage(content=response_text) if AIMessage else response_text], "response": response_text}


def resume_node(state):
    query = state["messages"][-1].content if hasattr(state["messages"][-1], "content") else str(state["messages"][-1])
    resume_data = state.get("resume_data", {}) or {}
    result = build_resume(llm, {"resume_data": resume_data, "query": query})
    suggestions_formatted = (
        "\n".join(["• " + s for s in result["suggestions"]])
        if result["suggestions"]
        else "• Resume details look solid!"
    )
    message = f"""
ATS Score: {result['ats']} / 100

Suggestions:
{suggestions_formatted}

Resume Content:
{result['resume']}
"""
    return {"messages": [AIMessage(content=message) if AIMessage else message], "response": message, "resume_data": result}


def general_node(state):
    query = state["messages"][-1].content if hasattr(state["messages"][-1], "content") else str(state["messages"][-1])
    search_result = web_search(query)
    if llm is None or GENERAL_PROMPT is None:
        response_text = f"Playure AI Sports Coach response for '{query}':\n\n{search_result if search_result else 'Focus on consistent training, balanced nutrition, and structured recovery routines tailored to your sport!'}"
        return {"messages": [AIMessage(content=response_text) if AIMessage else response_text], "response": response_text}

    try:
        response = (GENERAL_PROMPT | llm).invoke({"query": query, "search_result": search_result})
        return {"messages": [AIMessage(content=response.content) if AIMessage else response.content], "response": response.content}
    except Exception:
        response_text = f"Playure AI Sports Coach response for '{query}':\n\n{search_result if search_result else 'Focus on consistent training, balanced nutrition, and structured recovery routines tailored to your sport!'}"
        return {"messages": [AIMessage(content=response_text) if AIMessage else response_text], "response": response_text}


def router(state: SportsState):
    category = (state.get("category") or "general").lower()
    if category in ["physio", "injury", "stretch"]:
        return "physio"
    if category in ["training", "workout"]:
        return "training"
    if category in ["nutrition", "diet"]:
        return "nutrition"
    if category in ["resume", "cv"]:
        return "resume"
    return "general"


def extractor_node(state: SportsState):
    query = state["messages"][-1].content if hasattr(state["messages"][-1], "content") else str(state["messages"][-1])
    extracted = extract_information(query)
    existing = state.get("user_data", {})
    merged = {**existing, **{k: v for k, v in extracted.items() if v is not None}}
    return {"user_data": merged}


graph = None
if StateGraph is not None and START is not None and END is not None:
    try:
        builder = StateGraph(SportsState)
        builder.add_node("physio", physio_node)
        builder.add_node("training", training_node)
        builder.add_node("nutrition", nutrition_node)
        builder.add_node("resume", resume_node)
        builder.add_node("general", general_node)
        builder.add_node("extractor", extractor_node)

        builder.add_edge(START, "extractor")
        builder.add_conditional_edges(
            "extractor",
            router,
            {
                "physio": "physio",
                "training": "training",
                "nutrition": "nutrition",
                "resume": "resume",
                "general": "general",
            },
        )
        builder.add_edge("physio", END)
        builder.add_edge("training", END)
        builder.add_edge("nutrition", END)
        builder.add_edge("resume", END)
        builder.add_edge("general", END)

        graph = builder.compile()
    except Exception:
        graph = None


def process_ai_sports_query(
    query: str,
    category: Optional[str] = None,
    user_data: Optional[dict] = None,
    resume_data: Optional[dict] = None,
) -> Dict[str, Any]:
    # Auto-categorize based on text keywords if category is general or empty
    if not category or category.lower() == "general":
        q_lower = query.lower()
        if any(w in q_lower for w in ["pain", "injury", "stretch", "knee", "shoulder", "back", "ankle", "sprain", "physio"]):
            category = "physio"
        elif any(w in q_lower for w in ["workout", "training", "gym", "reps", "squat", "bench", "schedule", "overload", "1rm", "bowling", "footwork"]):
            category = "training"
        elif any(w in q_lower for w in ["diet", "nutrition", "protein", "calorie", "tdee", "bmr", "meal", "water", "food", "carbs"]):
            category = "nutrition"
        elif any(w in q_lower for w in ["resume", "cv", "ats", "scout"]):
            category = "resume"
        else:
            category = "general"

    state = {
        "messages": [HumanMessage(content=query)] if HumanMessage else [query],
        "category": category,
        "response": "",
        "user_data": user_data or {},
        "resume_data": resume_data or {},
        "tool_output": {},
    }

    if graph is not None:
        try:
            result = graph.invoke(state)
            return {
                "status": "success",
                "reply": result.get("response", ""),
                "category": category,
                "user_data": result.get("user_data", {}),
                "resume_data": result.get("resume_data", {}),
            }
        except Exception:
            pass

    # Direct execution fallback
    route_target = router(state)
    if route_target == "physio":
        res = physio_node(state)
    elif route_target == "training":
        res = training_node(state)
    elif route_target == "nutrition":
        res = nutrition_node(state)
    elif route_target == "resume":
        res = resume_node(state)
    else:
        res = general_node(state)

    return {
        "status": "success",
        "reply": res.get("response", ""),
        "category": category,
        "user_data": state.get("user_data", {}),
        "resume_data": res.get("resume_data", {}),
    }
